from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

try:
    from PIL import Image
except ImportError:  # pragma: no cover - Pillow is available in this repo, but keep a fallback.
    Image = None


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
DOCX_PATH = OUTPUT_DIR / "chinese_truth.docx"
MD_PATH = OUTPUT_DIR / "chinese_truth.md"
IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>.*)\]\((?P<path>.+)\)$")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for style_name, bold in [
        ("Title", True),
        ("Heading 1", True),
        ("Heading 2", True),
        ("Heading 3", True),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(12)
        style.font.bold = bold

    if "CodeBlock" not in document.styles:
        code_style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(12)
        code_pf = code_style.paragraph_format
        code_pf.line_spacing = 1.0
        code_pf.space_after = Pt(0)
        code_pf.space_before = Pt(0)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def flush_paragraph(document: Document, lines: list[str]) -> None:
    if lines:
        document.add_paragraph(" ".join(line.strip() for line in lines))
        lines.clear()


def flush_code(document: Document, lines: list[str]) -> None:
    if lines:
        paragraph = document.add_paragraph(style="CodeBlock")
        paragraph.add_run("\n".join(lines))
        lines.clear()


def resolve_markdown_path(raw_path: str) -> Path:
    candidate = raw_path.strip()
    if candidate.startswith("<") and candidate.endswith(">"):
        candidate = candidate[1:-1]
    path = Path(candidate)
    if not path.is_absolute():
        path = (MD_PATH.parent / path).resolve()
    return path


def add_markdown_image(document: Document, alt: str, raw_path: str) -> None:
    path = resolve_markdown_path(raw_path)
    if not path.exists():
        missing = document.add_paragraph(f"[Missing image] {path}")
        missing.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    max_width_inches = 165 / 25.4
    width = Mm(165)
    if Image is not None:
        with Image.open(path) as img:
            dpi_x = img.info.get("dpi", (96, 96))[0] or 96
            if not 20 <= dpi_x <= 1200:
                dpi_x = 96
            native_width_inches = img.size[0] / dpi_x
        width = Inches(min(max_width_inches, native_width_inches))

    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(path), width=width)

    if alt.strip():
        caption = document.add_paragraph(alt.strip())
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_markdown(document: Document, markdown: str) -> None:
    paragraph_lines: list[str] = []
    code_lines: list[str] = []
    in_code = False

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip("\n")

        if line.startswith("```"):
            flush_paragraph(document, paragraph_lines)
            if in_code:
                flush_code(document, code_lines)
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_paragraph(document, paragraph_lines)
            continue

        image_match = IMAGE_PATTERN.match(line.strip())
        if image_match:
            flush_paragraph(document, paragraph_lines)
            add_markdown_image(document, image_match.group("alt"), image_match.group("path"))
            continue

        if line.startswith("# "):
            flush_paragraph(document, paragraph_lines)
            title = document.add_paragraph(style="Title")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.add_run(line[2:].strip())
            continue

        heading_match = re.match(r"^(##|###|####) (.+)$", line)
        if heading_match:
            flush_paragraph(document, paragraph_lines)
            level = {"##": 1, "###": 2, "####": 3}[heading_match.group(1)]
            document.add_heading(heading_match.group(2).strip(), level=level)
            continue

        if re.match(r"^- ", line):
            flush_paragraph(document, paragraph_lines)
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\. ", line):
            flush_paragraph(document, paragraph_lines)
            document.add_paragraph(re.sub(r"^\d+\. ", "", line).strip(), style="List Number")
            continue

        paragraph_lines.append(line)

    flush_paragraph(document, paragraph_lines)
    flush_code(document, code_lines)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = MD_PATH.read_text(encoding="utf-8")

    document = Document()
    configure(document)
    render_markdown(document, markdown)
    document.save(str(DOCX_PATH))

    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
